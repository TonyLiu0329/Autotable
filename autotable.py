import os
import pandas as pd
import logging
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import json
from datetime import datetime

import re

logger = logging.getLogger(__name__)

class AutoTable:
    """自动化填表处理核心类"""
    def __init__(self, knowledge_base_path, word_template_path, llm_client, output_folder="output"):
        self.knowledge_base_path = knowledge_base_path
        self.word_template_path = word_template_path
        self.output_folder = output_folder
        self.llm_client = llm_client
        self.knowledge_base = None
        self.knowledge_dict = None
        self.doc = None

        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

    def load_knowledge_base(self):
        try:
            logger.info(f"正在加载知识库: {self.knowledge_base_path}")
            if self.knowledge_base_path.endswith('.xlsx'):
                # 读取所有工作表，不将第一行作为表头（header=None）
                # 这样可以处理 Key-Value 型的表格，也可以避免错误列名的问题
                dfs = pd.read_excel(self.knowledge_base_path, sheet_name=None, header=None)
                
                # 构建结构化知识库：{ "Sheet名": [ [行1数据], [行2数据], ... ], ... }
                structured_data = {}
                total_records = 0
                
                for sheet_name, df in dfs.items():
                    # 处理 NaN 值为 None/空字符串
                    df = df.where(pd.notnull(df), "")
                    
                    # 转换为二维列表 (List of Lists)
                    # 这种格式最通用，既适合列表型表格，也适合 KV 型表格
                    # LLM 可以根据数据分布自行推断行列关系
                    matrix_data = df.values.tolist()
                    
                    # 特殊处理 Text_Content: 如果它是原来的格式（有 Header），pd.read_excel(header=None) 会把 Header 也读成第一行数据
                    # 但 Text_Content 比较简单，即使把 'Content', 'Type' 当作数据也不影响理解
                    
                    structured_data[sheet_name] = matrix_data
                    total_records += len(matrix_data)
                
                self.knowledge_dict = structured_data
                logger.info(f"Excel知识库加载完成，共读取 {len(dfs)} 个工作表，合计 {total_records} 条数据")
                
                # 输出结构化后的字典供调试/查看
                logger.info(f"结构化处理后的知识库字典: {json.dumps(self.knowledge_dict, ensure_ascii=False, default=str)}")
                
                return True
            elif self.knowledge_base_path.endswith('.json'):
                # 直接读取 JSON 格式的知识库
                with open(self.knowledge_base_path, 'r', encoding='utf-8') as f:
                    self.knowledge_dict = json.load(f)
                logger.info(f"JSON知识库加载完成: {self.knowledge_base_path}")
                return True
            else:
                logger.error("目前只支持 .xlsx 或 .json 格式知识库")
                return False
        except Exception as e:
            logger.error(f"知识库加载异常: {str(e)}")
            return False

    def load_template(self):
        try:
            logger.info(f"正在加载Word模板: {self.word_template_path}")
            self.doc = Document(self.word_template_path)
            logger.info(f"模板加载完成，包含{len(self.doc.tables)}个表格")
            return True
        except Exception as e:
            error_msg = str(e)
            if "no relationship of type" in error_msg or "File is not a zip file" in error_msg:
                logger.error(f"模板加载失败: 文件格式不正确。请确保您上传的是真正的 .docx 文件，而不是直接修改后缀名的 .doc 文件。请尝试用 Word 打开该文件并'另存为' .docx 格式。({error_msg})")
            else:
                logger.error(f"模板加载失败: {error_msg}")
            return False

    def _is_potential_slot(self, text):
        """判断文本是否为填空位"""
        if not text:
            return True
            
        # 移除常见的不可见字符和空白
        clean_text = text.strip().replace('\u200b', '').replace('\u3000', ' ')
        if not clean_text:
            return True
            
        # 负向规则：排除明显的表头/Label（防止误判）
        # 1. "第( )完成人"、"第( )完成单位" 类型的表头
        if re.match(r'^第\s*[（(]\s*[）)]\s*(?:完成人|作者|完成单位|单位|起草人)', clean_text):
            return False
            
        if set(clean_text).issubset(set(" _()（）")):
            return True
        
        # 正则增强规则
        # 包含连续下划线
        if re.search(r'[_]{2,}', clean_text):
            return True
        # 包含空括号 或 提示性括号
        elif re.search(r'[(\uff08](?:\s*|.*?(?:填写|输入|粘贴|限|字|内容).*?)[)\uff09]', clean_text):
            return True
        # 包含 "年" 和 "月" 的日期格式
        elif '年' in clean_text and '月' in clean_text:
            if not re.match(r'^\d', clean_text):
                 if re.search(r'[_]+|\s+', clean_text):
                     return True
        # 启发式规则：以冒号结尾的 Prompt
        elif re.search(r'[:：]\s*$', clean_text):
            return True
        # 针对大表格的长文本Prompt（如 "1. 成果简介..."）
        # 特征：以数字序号开头，且长度超过一定阈值，暗示这是一个问题描述而非简单标题
        elif re.match(r'^\d+[.、\s]', clean_text) and len(clean_text) > 5:
            return True
            
        return False

    def _has_visual_placeholder(self, element):
        """
        检查段落或单元格是否包含“视觉占位符”（带下划线的空白区域）。
        """
        paragraphs = []
        if hasattr(element, 'paragraphs'):
            paragraphs = element.paragraphs
        else:
            paragraphs = [element] # Assume it's a paragraph
            
        for para in paragraphs:
            for run in para.runs:
                # 检查是否有下划线
                is_underlined = run.underline is not None and run.underline is not False
                # 检查内容是否主要为空白
                text = run.text
                is_blank = all(c in ' \t\u3000\u00A0' for c in text)
                
                if is_underlined and is_blank and len(text) >= 2:
                    return True
        return False

    def _preprocess_paragraphs(self, paragraphs):
        """
        预处理段落列表：识别填空位
        返回: (markdown_text, anchor_map, id_to_text_map)
        anchor_map: {anchor_id: paragraph_index}
        """
        anchor_map = {}
        id_to_text_map = {}
        markdown_lines = []
        anchor_counter = 1
        
        for idx, para in enumerate(paragraphs):
            text = para.text.strip()
            
            # 检查是否包含视觉占位符（如下划线空格）
            has_visual_slot = self._has_visual_placeholder(para)
            
            # 修复：跳过空段落，除非它包含视觉占位符
            if not text and not has_visual_slot:
                continue
                
            if self._is_potential_slot(text) or has_visual_slot:
                anchor_id = f"{{{{ID_{anchor_counter:03d}}}}}"
                anchor_map[anchor_id] = idx
                id_to_text_map[anchor_id] = f"原内容: '{text}'"
                
                # 在Markdown中展示
                # 尝试更智能的展示：如果段落很短，直接展示ID；如果长，展示部分上下文？
                # 暂时简单处理：直接用 ID 替换原文本展示给 LLM
                display_text = text if text else "[下划线填空区]"
                markdown_lines.append(f"- {anchor_id} (原: {display_text})")
                anchor_counter += 1
            else:
                # 如果不是填空位，但也包含在文档中，是否给LLM看？
                # 如果是纯文本，可能包含上下文信息。
                # 最好还是提供一些上下文。
                if len(text) > 0:
                    markdown_lines.append(f"- {text}")

        markdown_text = "\n".join(markdown_lines)
        return markdown_text, anchor_map, id_to_text_map

    def _preprocess_table(self, table):
        """
        预处理表格：识别填空位，生成带锚点的Markdown文本，并记录锚点映射。
        返回: (markdown_text, anchor_map, id_to_text_map)
        """
        anchor_map = {}
        id_to_text_map = {}
        markdown_lines = []
        anchor_counter = 1
        
        # 使用列表存储已处理的单元格 _tc 对象，以处理合并单元格
        processed_tcs = [] # Store (tc_object, anchor_id)
        
        # 遍历每一行
        for row_idx, row in enumerate(table.rows):
            row_cells_text = []
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 检查该单元格是否已处理过
                current_tc = cell._tc
                existing_anchor_id = None
                for seen_tc, seen_id in processed_tcs:
                    if current_tc == seen_tc:
                        existing_anchor_id = seen_id
                        break
                
                if existing_anchor_id:
                    # 如果已处理过，直接使用之前的 ID
                    row_cells_text.append(existing_anchor_id)
                    # 注意：我们不需要再次添加到 anchor_map，因为之前已经加过了
                    # 但我们需要确保在 Markdown 中显示这个 ID，以便 LLM 理解表格结构
                    continue

                # 判断是否是填空位
                is_potential_slot = self._is_potential_slot(cell_text)
                has_visual_slot = self._has_visual_placeholder(cell)
                
                if is_potential_slot or has_visual_slot:
                    anchor_id = f"{{{{ID_{anchor_counter:03d}}}}}"
                    anchor_map[anchor_id] = (row_idx, col_idx)
                    
                    # 记录为已处理
                    processed_tcs.append((current_tc, anchor_id))
                    
                    # 尝试获取上下文提示（Label）
                    context_hint = ""
                    try:
                        # 尝试获取左侧单元格文本作为提示
                        if col_idx > 0:
                            left_text = row.cells[col_idx-1].text.strip()
                            if left_text and len(left_text) < 20:
                                context_hint = f" | 左侧Label: {left_text}"
                        
                        # 如果左侧为空，尝试获取上方单元格文本（针对上下结构的表格）
                        if not context_hint and row_idx > 0:
                            top_text = table.cell(row_idx-1, col_idx).text.strip()
                            if top_text and len(top_text) < 20:
                                context_hint = f" | 上方Label: {top_text}"
                    except Exception:
                        pass

                    # 记录原始文本和上下文提示，供LLM参考
                    id_to_text_map[anchor_id] = f"原内容: '{cell_text}'{context_hint}"
                    row_cells_text.append(anchor_id)
                    anchor_counter += 1
                else:
                    # 调试日志：记录为什么这个单元格没有被选中（采样打印）
                    # if len(cell_text) > 0 and len(cell_text) < 10:
                    #    logger.debug(f"跳过单元格: '{cell_text}' - 未匹配规则")
                    
                    # 记录非填空位单元格为已处理（虽然没有ID，但我们不想重复处理）
                    # 不过，非填空位通常有内容，如果重复出现，在Markdown里重复显示内容是正确的
                    # 所以这里不需要记录 processed_tcs，除非我们想去重显示？
                    # 对于Markdown表格，如果合并单元格跨列，我们通常希望显示 | Content | Content | Content |
                    # 这样结构是对齐的。所以这里保持原样。
                    
                    # 清理一下换行符，以免破坏Markdown表格结构
                    clean_text = cell_text.replace('\n', '<br>')
                    row_cells_text.append(clean_text)
            
            markdown_lines.append("| " + " | ".join(row_cells_text) + " |")
            
        # 组合成Markdown表格字符串
        markdown_text = "\n".join(markdown_lines)
        return markdown_text, anchor_map, id_to_text_map

    def analyze_tables_with_llm(self, table_markdown, knowledge_context, id_to_text_map, used_contexts=None):
        # 动态判断知识库格式，生成不同的 Prompt 描述
        data_format_desc = "扁平化的 JSON 键值对（Key-Value）" if isinstance(knowledge_context, dict) else "按 Sheet（来源）分组的二维数组（矩阵）格式"
        
        used_context_desc = ""
        if used_contexts and len(used_contexts) > 0:
            used_context_desc = f"""
        **上下文去重约束（重要）**：
        当前文档中包含多个结构相似的表格，用于填写不同实体（人员/项目）的信息。
        以下实体标识（如姓名、项目名）**已被前面的表格使用过**：
        {json.dumps(used_contexts, ensure_ascii=False)}
        
        **请务必从知识库中选择一个【未使用过】的新实体数据进行填充。**
        - 如果知识库是人员列表，请选择下一个不同的人员。
        - 如果知识库是项目列表，请选择下一个不同的项目。
        - 如果确实没有更多新数据，才允许重复。
        """

        prompt = f"""
        请分析以下带有锚点（格式如 {{{{ID_XXX}}}}）的文档内容（表格或段落），并结合提供的知识库数据，将正确的值填入对应的锚点。
        
        知识库数据（{data_format_desc}）：
        {json.dumps(knowledge_context, ensure_ascii=False, default=str)}
        
        文档内容结构（Markdown）：
        {table_markdown}
        
        锚点对应的原始文本（参考用，可能包含提示信息）：
        {json.dumps(id_to_text_map, ensure_ascii=False)}
        {used_context_desc}
        
        **核心原则：严格基于知识库**
        1. **绝对禁止编造数据**：你只能使用“知识库数据”中显式提供的信息。
        3. **禁止推测**：不要根据常识或上下文去猜测缺失的信息（例如：不要自己编造邮编、电话、日期，也不要推测上级单位）。原文没有就是没有。

        请仔细思考字段的对应关系。
        注意：知识库数据可能为以下两种格式之一：
        1. 按 Sheet（来源）分组的二维数组（矩阵）格式。
        2. 扁平化的 JSON 键值对（Key-Value）。
        
        请根据上下文自动推断。
        - 如果遇到包含换行符被拼接的字段名（如“现从事工作及专长”），请尝试模糊匹配。
        - **必须**参考提供的“锚点对应的原始文本”中的上下文线索（如“左侧Label”或“上方Label”）来确定填入内容。
        
        特别注意以下字段的提取和填充：
        1. **多键合并与分发**：
           - **合并**：如果表格中只有一个对应字段的单元格，但知识库中有多个相关键（如 "奖项_1", "奖项_2"），请将它们合并为一个完整段落（如 "1. A奖；2. B奖"）。
           - **分发**：如果表格中针对同一属性（如“爱好”）预留了**多个独立**的单元格（即多个锚点），且知识库中有对应的多条数据，请将数据**分散填入**不同的锚点，不要重复。
              - 例如：表格“爱好”下有 `{{{{ID_001}}}}` 和 `{{{{ID_002}}}}`，知识库有 `爱好_1: A`, `爱好_2: B`。
              - 正确：`"{{{{ID_001}}}}": "A"`, `"{{{{ID_002}}}}": "B"`
              - 错误：`"{{{{ID_001}}}}": "A, B"`, `"{{{{ID_002}}}}": "A, B"`
              - **注意**：如果数据条数少于单元格数（如只有 `爱好_1: A`），请只填入第一个锚点，其余锚点**不要**在返回的JSON中出现（即留空）。
        2. **复杂文本字段**：如“现从事工作及专长”、“何时何地受何奖励”。这些内容可能较长，包含多行，请务必提取完整。
        3. **基础信息字段**：如“工作单位”、“电子信箱”、“通讯地址”。这些信息可能散落在不同位置，请仔细查找。
        4. **格式处理**：如果源数据中包含换行符（\\n），请根据目标表格的语境合理保留或替换为逗号/空格。
        5. **问答式填充**：
           - 如果锚点对应的“原始文本”是一个问题或指令（例如 "1. 成果简介..."），请**只返回该问题的答案内容**，不要重复问题本身。程序会自动将答案追加在问题下方。

        如果一个字段在多个地方出现，请优先选择最匹配上下文的值。
        
        **禁止行为**：
        - 严禁在表格末尾或其他空位自动生成“总结”、“备注”或“额外说明”，除非表格中有明确的“备注”或“总结”标签指示。
        - 如果锚点没有明确的上下文指示（Label），且无法确定其对应关系，请保持为空，不要强行填入剩余的知识库信息。
        - **再次强调**：知识库中不存在的信息，填入值必须为空字符串。

        重要：请返回**完整**的填入内容。
        如果原始文本是占位符（如 "____"），直接返回填入值。
        如果原始文本包含提示信息且你需要保留（例如 "姓名：____"），请返回完整内容（例如 "姓名：张三"）。
        如果原始文本是日期格式（如 "____年__月__日"），请返回填充好的完整日期字符串（例如 "2024年1月1日"）。
        通常情况下，对于包含下划线的单元格，用户希望你填充内容并覆盖原有占位符。
        
        返回 JSON 格式：
        {{
            "__identity__": "这里填入你本次使用的实体唯一标识（如姓名：张三），用于后续去重",
            "{{{{ID_001}}}}": "填入的值1",
            "{{{{ID_002}}}}": "",
            ...
        }}
        请确保只返回JSON格式数据，不要包含其他内容。
        """
        try:
            messages = [
                {"role": "system", "content": "你是一个专业的文档填充助手，擅长处理复杂表格和多层级数据映射。"},
                {"role": "user", "content": prompt}
            ]
            result = self.llm_client.chat_completion(messages, temperature=0.1)
            json_str = self._extract_json(result)
            return json.loads(json_str)
        except Exception as e:
            logger.error(f"表格分析失败: {str(e)}")
            return {}

    def _extract_json(self, text):
        try:
            json.loads(text)
            return text
        except json.JSONDecodeError:
            import re
            # 移除 replace('\n', '') 以保留 JSON 字符串中的换行符
            json_match = re.search(r'({.*})', text, re.DOTALL)
            if json_match:
                return json_match.group(1)
            raise ValueError("未找到有效JSON内容")

    def _extract_paragraph_char_style(self, paragraph):
        """从段落属性(pPr)中提取默认字符样式"""
        style = {}
        try:
            # 访问底层 XML 元素
            p = paragraph._element
            if p.pPr is None:
                return style
            
            # 安全获取 rPr (Run Properties)
            # 注意：pPr 是一个 CT_PPr 对象，它可能没有直接的 .rPr 属性访问器
            # 我们应该使用 find 方法来查找子元素
            rPr = p.pPr.find(qn('w:rPr'))
            
            if rPr is None:
                return style
            
            # 1. 字体名称
            # rPr.rFonts 可能也是通过 find 获取
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                # 优先取中文字体(eastAsia)，其次 ascii
                font_name = rFonts.get(qn('w:eastAsia')) or rFonts.get(qn('w:ascii'))
                if font_name:
                    style['name'] = font_name
            
            # 2. 字号 (XML中是半点，1/144英寸)
            sz = rPr.find(qn('w:sz'))
            if sz is not None and sz.val is not None:
                try:
                    # Pt(1) = 2 half-points
                    # 正常情况：half-points / 2 = points
                    # 异常情况：某些文档中可能存储的是 EMU 值 (1 pt = 12700 EMU)
                    val = int(sz.val)
                    
                    # 阈值判断：如果值大于 4000 (即 2000pt)，几乎可以肯定是 EMU 单位
                    # 正常的字号通常在 1-100pt (2-200 half-points) 之间
                    if val > 4000:
                        style['size'] = val # 直接作为 EMU 使用
                    else:
                        style['size'] = Pt(val / 2) # 作为 half-points 转换
                except Exception:
                    pass

            # 3. 加粗
            b = rPr.find(qn('w:b'))
            if b is not None:
                # 标签存在即为真，除非显式设为 false/0
                val = b.val
                style['bold'] = False if val in ['0', 'false', 'off'] else True
            
            # 4. 颜色
            color = rPr.find(qn('w:color'))
            if color is not None and color.val is not None:
                hex_color = color.val
                if hex_color != 'auto':
                    try:
                        style['color'] = RGBColor.from_string(hex_color)
                    except Exception:
                        pass
                        
            # 5. 斜体
            i = rPr.find(qn('w:i'))
            if i is not None:
                 val = i.val
                 style['italic'] = False if val in ['0', 'false', 'off'] else True
                 
        except Exception as e:
            logger.warning(f"提取段落默认样式失败: {e}")
            
        return style

    def _extract_run_style(self, run):
        """提取Run的字体样式"""
        style = {}
        if not run:
            return style
        
        # 基础属性
        if run.font.name:
            style['name'] = run.font.name
        if run.font.size:
            style['size'] = run.font.size
        if run.font.bold is not None:
            style['bold'] = run.font.bold
        if run.font.italic is not None:
            style['italic'] = run.font.italic
        if run.font.color and run.font.color.rgb:
            style['color'] = run.font.color.rgb
        if run.font.underline:
            style['underline'] = run.font.underline
            
        if style:
            logger.info(f"成功提取字体样式: {style}")
        else:
            logger.info("未提取到显式字体样式 (可能是默认样式，将继承段落设置)")
            
        return style

    def _apply_run_style(self, run, style):
        """应用字体样式到Run"""
        if not style:
            return
            
        if 'name' in style:
            run.font.name = style['name']
            # 设置中文字体
            try:
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), style['name'])
            except Exception as e:
                logger.warning(f"设置中文字体失败: {e}")
            
        if 'size' in style:
            run.font.size = style['size']
        if 'bold' in style:
            run.font.bold = style['bold']
        if 'italic' in style:
            run.font.italic = style['italic']
        if 'color' in style:
            run.font.color.rgb = style['color']
        if 'underline' in style:
            run.font.underline = style['underline']

    def _smart_fill_paragraph(self, para, value):
        """
        尝试智能填充：如果段落包含下划线/空格占位符，则只替换占位符部分，并保留下划线格式。
        返回 True 表示已处理，False 表示未匹配到占位符，需调用方回退到默认逻辑。
        """
        runs = para.runs
        placeholder_idx = -1
        
        # 1. 寻找段落末尾的占位符 Run
        # 占位符特征：有下划线，且内容主要是空格、下划线、制表符
        for i, run in enumerate(runs):
            text = run.text
            is_underlined = run.underline is not None and run.underline is not False
            is_placeholder_chars = all(c in ' _\t\u3000\u00A0' for c in text)
            
            # 必须有一定的长度（避免误判单个空格），或者是纯下划线
            if is_underlined and (len(text) >= 1 and is_placeholder_chars):
                # 进一步检查：如果是空格，必须是下划线的空格。
                # 这里假设 run.underline 已经过滤了无下划线的情况
                placeholder_idx = i
                break
        
        if placeholder_idx == -1:
            return False

        # 2. 提取 Label 和 Value
        # Label 是占位符之前的所有文本
        label_text = "".join([r.text for r in runs[:placeholder_idx]])
        
        # 归一化处理以进行模糊匹配
        import re
        label_clean = re.sub(r'\s+', '', label_text)
        value_str = str(value)
        value_clean = re.sub(r'\s+', '', value_str)
        
        fill_content = ""
        
        # 尝试从 value_str 中剥离 label
        # 情况A: value_str 包含 label (例如 "姓名: 张三")
        # 我们需要在 value_str 中找到 label_clean 的结束位置
        
        # 简单的字符串包含检查
        if label_clean and label_clean in value_clean:
            # 找到 label 在 value 中的位置
            # 这比较难精确对应到 value_str 的索引，因为空格差异。
            # 采用字符逐个匹配法
            val_ptr = 0
            lbl_ptr = 0
            match_end_idx = 0
            
            while val_ptr < len(value_str) and lbl_ptr < len(label_clean):
                if value_str[val_ptr].isspace():
                    val_ptr += 1
                    continue
                
                if value_str[val_ptr] == label_clean[lbl_ptr]:
                    val_ptr += 1
                    lbl_ptr += 1
                    match_end_idx = val_ptr
                else:
                    break
            
            if lbl_ptr == len(label_clean):
                fill_content = value_str[match_end_idx:].strip()
            else:
                # 匹配失败，可能 LLM 修改了 Label
                # 这种情况下，为了安全，我们假设整个 value_str 都是内容？
                # 或者回退到 False?
                # 如果回退，会覆盖 Label。如果填入，会重复 Label。
                # 优先保护 Label 不被覆盖。
                return False
        else:
            # 情况B: value_str 不包含 label (例如 LLM 只返回了 "张三")
            # 直接把 value_str 当作填充内容
            fill_content = value_str.strip()

        # 3. 执行填充
        # 更新占位符 Run
        target_run = runs[placeholder_idx]
        # 在内容前后加空格以保持美观（可选，视模板而定，这里加一个前导空格防止紧贴）
        # 只有当 fill_content 不为空时才填充，否则保持原样（或者清空？）
        # 用户通常希望填入内容。
        target_run.text = " " + fill_content + " " 
        target_run.underline = True # 强制下划线
        
        # 清除后续的占位符 Run (防止原占位符很长，被分成了多段)
        for i in range(placeholder_idx + 1, len(runs)):
            r = runs[i]
            is_underlined = r.underline is not None and r.underline is not False
            is_placeholder_chars = all(c in ' _\t\u3000\u00A0' for c in r.text)
            if is_underlined and is_placeholder_chars:
                r.text = ""
            else:
                # 遇到非占位符（如后面的括号说明），停止清除
                break
                
        return True

    def fill_document(self):
        if not self.doc or self.knowledge_dict is None:
            logger.error("文档或知识库未正确初始化")
            return False
        
        filled_count = 0
        used_identities = [] # 全局追踪已使用的实体标识
        
        # --- 1. 处理正文段落 ---
        logger.info("正在处理正文段落...")
        para_markdown, para_anchor_map, para_id_to_text_map = self._preprocess_paragraphs(self.doc.paragraphs)
        
        if para_anchor_map:
            logger.info(f"发现 {len(para_anchor_map)} 个段落填空位，正在请求 LLM 分析...")
            fill_map = self.analyze_tables_with_llm(para_markdown, self.knowledge_dict, para_id_to_text_map, used_contexts=used_identities)
            
            # 提取并记录本次使用的实体标识
            identity = fill_map.pop("__identity__", None)
            if identity:
                used_identities.append(identity)
                logger.info(f"正文段落使用了实体: {identity}")
            
            for anchor_id, value in fill_map.items():
                if anchor_id in para_anchor_map:
                    try:
                        idx = para_anchor_map[anchor_id]
                        para = self.doc.paragraphs[idx]
                        
                        # 智能填充逻辑
                        original_text = para.text.strip()
                        clean_val = str(value).strip()

                        # 优先尝试智能下划线填充
                        if self._smart_fill_paragraph(para, clean_val):
                            filled_count += 1
                            logger.debug(f"段落锚点 {anchor_id} 采用下划线保留模式填充")
                            continue
                        
                        # 检查是否是 "Label: " 形式的纯标签段落（无下划线占位符）
                        # 如果是，且 LLM 返回的值没有包含 Label，则采用追加模式，防止覆盖标签
                        is_pure_label = re.search(r'[:：]\s*$', original_text)
                        
                        if is_pure_label and not clean_val.startswith(original_text[:min(5, len(original_text))]):
                             # 追加模式：保留原标签，追加内容
                             para.add_run(f" {clean_val}")
                             logger.debug(f"段落锚点 {anchor_id} 采用追加模式填充")
                        else:
                            # 覆盖模式：全段替换
                            # 注意：这会丢失段落内的部分格式（如加粗），但保留段落整体样式
                            style = para.style
                            para.text = clean_val
                            para.style = style
                        
                        filled_count += 1
                        logger.debug(f"段落锚点 {anchor_id} 已填充值: {value}")
                    except Exception as e:
                        logger.error(f"段落填充异常: {anchor_id} - {str(e)}")
                else:
                    logger.warning(f"LLM返回了不存在的段落锚点ID: {anchor_id}")
        else:
            logger.info("正文段落中未发现填空位")

        # --- 2. 处理表格 ---
        for table_idx, table in enumerate(self.doc.tables):
            logger.info(f"正在处理第{table_idx + 1}个表格")
            
            # 第一阶段：Word 模板的“数字化”预处理
            table_markdown, anchor_map, id_to_text_map = self._preprocess_table(table)
            
            if not anchor_map:
                logger.info("该表格未发现填空位，跳过")
                continue

            # 第二阶段：LLM 分析并直接返回填充映射
            # 将知识库数据作为上下文传给 LLM，并传入已使用的实体列表
            fill_map = self.analyze_tables_with_llm(table_markdown, self.knowledge_dict, id_to_text_map, used_contexts=used_identities)
            
            # 提取并记录本次使用的实体标识
            identity = fill_map.pop("__identity__", None)
            if identity:
                used_identities.append(identity)
                logger.info(f"表格 {table_idx + 1} 使用了实体: {identity}")
            
            for anchor_id, value in fill_map.items():
                if anchor_id in anchor_map:
                    try:
                        row, col = anchor_map[anchor_id]
                        cell = table.cell(row, col)
                        
                        # 检查是否需要追加模式 (Append Mode)
                        # 如果原单元格包含类似 "1. xxx (不超过xx字)" 的指令性文本，且 LLM 返回的内容不包含该头信息，则追加
                        original_text = cell.text.strip()
                        should_append = False
                        
                        # 识别指令性表头特征：
                        # 1. 以数字开头 (1. / 1、 / 1 )
                        # 2. 包含字数限制说明 (不超过...字)
                        # 3. 新增：基于长度和数字开头的宽松匹配，与 _is_potential_slot 保持一致
                        if len(original_text) > 0 and (
                            re.match(r'^\d+(\.\d+)*[、. ]', original_text) or 
                            re.search(r'[（(].*?不超过.*?字.*?[)）]', original_text) or
                            (re.match(r'^\d+[.、\s]', original_text) and len(original_text) > 5)
                        ):
                            # 进一步检查：如果 LLM 返回的值已经包含了原文本（或者原文本的前一部分），则不需要追加，直接覆盖即可
                            # 简单的模糊检查
                            clean_val = str(value).strip()
                            if clean_val.startswith(original_text[:min(10, len(original_text))]):
                                should_append = False
                            else:
                                should_append = True
                        
                        if should_append:
                            logger.info(f"检测到指令性表头，采用追加模式: {anchor_id}")
                            
                            # 1. 尝试从现有标题（第一段）提取样式
                            font_style = {}
                            if cell.paragraphs and cell.paragraphs[0].runs:
                                font_style = self._extract_run_style(cell.paragraphs[0].runs[0])
                                logger.info(f"锚点 {anchor_id}: 追加模式下，成功从标题提取字体样式: {font_style}")
                            
                            # 2. 追加新段落
                            # 注意：add_paragraph 会在单元格末尾添加新段落
                            new_para = cell.add_paragraph(str(value))
                            
                            # 3. 应用样式
                            if font_style:
                                for run in new_para.runs:
                                    self._apply_run_style(run, font_style)
                        else:
                            # 尝试保留原有样式
                            filled_via_smart = False
                            if cell.paragraphs:
                                # 优先尝试智能下划线填充
                                if self._smart_fill_paragraph(cell.paragraphs[0], value):
                                    filled_via_smart = True
                            
                            if not filled_via_smart:
                                if cell.paragraphs:
                                    # 获取第一段
                                    first_para = cell.paragraphs[0]
                                    
                                    # 调试信息：检查Runs状态
                                    if not first_para.runs:
                                        logger.info(f"锚点 {anchor_id}: 单元格段落无 Runs (可能是空单元格)，无法提取预设样式。建议在模板中输入一个空格并设置样式。")
                                    
                                    # 尝试提取字体样式（从第一个Run）
                                    font_style = {}
                                    if first_para.runs:
                                        font_style = self._extract_run_style(first_para.runs[0])
                                    else:
                                        # 如果当前单元格为空（无Run），尝试从段落属性(pPr)中提取预设的字符样式
                                        # 这通常是用户在空单元格中设置的格式
                                        font_style = self._extract_paragraph_char_style(first_para)
                                        if font_style:
                                            logger.info(f"锚点 {anchor_id}: 成功从空单元格提取预设样式: {font_style}")
                                        else:
                                            logger.info(f"锚点 {anchor_id}: 当前单元格为空且无预设样式，将使用默认样式填充")

                                    # 清空段落内容但保留段落属性
                                    # 注意：first_para.clear() 会清除所有runs，但保留段落样式
                                    first_para.clear()
                                    
                                    # 添加新内容
                                    new_run = first_para.add_run(str(value))
                                    
                                    # 应用字体样式
                                    if font_style:
                                        self._apply_run_style(new_run, font_style)
                                else:
                                    # 确保值为字符串
                                    cell.text = str(value)
                            
                        filled_count += 1
                        logger.debug(f"表格锚点 {anchor_id} 已填充值: {value}")
                    except IndexError:
                        logger.error(f"无效的单元格位置: {anchor_map.get(anchor_id)}")
                    except Exception as e:
                        logger.error(f"字段填写异常: {anchor_id} - {str(e)}")
                else:
                    logger.warning(f"LLM返回了不存在的表格锚点ID: {anchor_id}")
                    
        logger.info(f"完成文档填充，共填写{filled_count}个字段")
        return True

    def save_document(self, filename=None):
        if not self.doc:
            logger.error("文档实例未初始化")
            return False
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = os.path.splitext(os.path.basename(self.word_template_path))[0]
            filename = f"{base_name}_filled_{timestamp}.docx"
        output_path = os.path.join(self.output_folder, filename)
        try:
            self.doc.save(output_path)
            logger.info(f"文档已保存至: {output_path}")
            return True
        except Exception as e:
            logger.error(f"文档保存失败: {str(e)}")
            return False

    def run(self):
        logger.info("启动自动化填表流程")
        if all([
            self.load_knowledge_base(),
            self.load_template(),
            self.fill_document(),
            self.save_document()
        ]):
            logger.info("流程执行成功")
            return True
        logger.error("流程执行过程中发生错误")
        return False