import os
import pandas as pd
from docx import Document
import logging

import re

import json

logger = logging.getLogger(__name__)

def extract_content_to_json(docx_path, output_json_path, llm_client):
    """
    使用 LLM 智能提取 Word 文档内容，生成扁平化的 JSON 知识库
    """
    try:
        try:
            doc = Document(docx_path)
        except Exception as e:
            error_msg = str(e)
            if "no relationship of type" in error_msg or "File is not a zip file" in error_msg:
                raise ValueError(f"文件格式不正确。请确保您上传的是真正的 .docx 文件，而不是直接修改后缀名的 .doc 文件。({error_msg})")
            raise e
            
        full_text = []

        # 1. 收集所有文本（页眉、页脚、正文、表格）
        # 页眉
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    full_text.append(f"[页眉] {para.text.strip()}")
            for para in section.footer.paragraphs:
                if para.text.strip():
                    full_text.append(f"[页脚] {para.text.strip()}")

        # 正文
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 表格
        for i, table in enumerate(doc.tables):
            full_text.append(f"\n[表格 {i+1}]")
            for row in table.rows:
                # 使用 ' // ' 替换换行符，以便在保持单行结构的同时保留换行信息
                row_text = " | ".join([cell.text.strip().replace('\n', ' // ') for cell in row.cells if cell.text.strip()])
                if row_text:
                    full_text.append(row_text)
        
        def _split_chunks(lines, max_chars=8000):
            chunks = []
            buf = []
            length = 0
            for line in lines:
                s = line if isinstance(line, str) else str(line)
                if length + len(s) + 1 > max_chars and buf:
                    chunks.append("\n".join(buf))
                    buf = []
                    length = 0
                buf.append(s)
                length += len(s) + 1
            if buf:
                chunks.append("\n".join(buf))
            return chunks

        def _parse_json(text):
            try:
                start = text.find('{')
                end = text.rfind('}') + 1
                if start != -1 and end != -1:
                    return json.loads(text[start:end])
                raise ValueError("未找到JSON内容")
            except Exception:
                raise

        def _merge(a, b):
            for k, v in b.items():
                if k not in a:
                    a[k] = v
                else:
                    av = a[k]
                    if isinstance(av, dict) and isinstance(v, dict):
                        _merge(av, v)
                    elif isinstance(av, list):
                        if isinstance(v, list):
                            a[k] = av + [x for x in v]
                        else:
                            a[k] = av + [v]
                    else:
                        if av == v:
                            a[k] = av
                        else:
                            a[k] = [av, v] if not isinstance(av, list) else av + [v]
            return a

        chunks = _split_chunks(full_text)
        total = len(chunks)
        merged = {}
        raw_fallback = {}

        for idx, chunk in enumerate(chunks, start=1):
            prompt = f"""
            请分析以下文档内容片段（第{idx}/{total}段），将其提取为扁平化 JSON 键值对。
            
            文档内容片段：
            {chunk}
            
            要求：
            1. 区分短事实与长段落，短事实细粒度拆分，长段落保留原文。
            2. 列表类内容可保留为序号文本或拆分为键_1、键_2。
            3. 键名简洁规范，特定小标题直接用原标题。
            4. 处理 ' // ' 作为换行提示，合理还原。
            5. 仅返回 JSON 对象。
            """
            messages = [
                {"role": "system", "content": "你是一个专业的数据提取专家，擅长将非结构化文档转化为结构化数据。"},
                {"role": "user", "content": prompt}
            ]
            try:
                response = llm_client.chat_completion(messages, temperature=0.1)
                data = _parse_json(response)
                merged = _merge(merged, data)
            except Exception as e:
                raw_fallback[f"Raw_Content_Chunk_{idx}"] = chunk
                raw_fallback[f"Error_Chunk_{idx}"] = str(e)

        final_data = merged if merged else {"Raw_Content": "\n".join(full_text), "Error": "JSON解析失败"}
        if raw_fallback:
            final_data.update(raw_fallback)

        with open(output_json_path, 'w', encoding='utf-8') as f:
            json.dump(final_data, f, ensure_ascii=False, indent=4)
        logger.info(f"成功使用 LLM 提取数据到 {output_json_path}")
        return True

    except Exception as e:
        logger.error(f"智能提取失败: {str(e)}")
        return False

def clean_cell_text(text):
    """
    清洗单元格文本：
    1. 去除首尾空白
    2. 尝试将混在一行的 Key-Value 对（如 '工作单位：xxx 职务：yyy'）拆分为多行
    """
    if not text:
        return ""
    
    text = text.strip()
    
    # 正则策略：查找 "空格 + 中文键名 + 冒号" 的模式，将其前面的空格替换为换行符
    # 键名限制为 2-6 个汉字，避免误伤普通句子
    # 处理中文冒号和英文冒号
    pattern = r'\s+([\u4e00-\u9fa5]{2,10}[：:])'
    text = re.sub(pattern, r'\n\1', text)
    
    return text

def extract_tables_from_docx(docx_path, output_excel_path):
    """
    从Word文档中提取表格和文本，保存为Excel文件
    """
    try:
        try:
            doc = Document(docx_path)
        except Exception as e:
            error_msg = str(e)
            if "no relationship of type" in error_msg or "File is not a zip file" in error_msg:
                raise ValueError(f"文件格式不正确。请确保您上传的是真正的 .docx 文件，而不是直接修改后缀名的 .doc 文件。({error_msg})")
            raise e

        # 1. 提取所有文本段落 (包括页眉页脚)
        text_content = []
        
        # 提取页眉页脚
        for section in doc.sections:
            # 页眉
            for para in section.header.paragraphs:
                 if para.text.strip():
                    text_content.append({"Content": para.text.strip(), "Type": "Header"})
            # 页脚
            for para in section.footer.paragraphs:
                 if para.text.strip():
                    text_content.append({"Content": para.text.strip(), "Type": "Footer"})

        # 正文段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append({"Content": para.text.strip(), "Type": "Text"})
        
        # 2. 提取表格
        tables_data = []
        for i, table in enumerate(doc.tables):
            # 尝试提取表格数据
            # 假设第一行是表头
            rows = []
            if len(table.rows) > 0:
                # 获取所有行数据
                for row in table.rows:
                    # 保留换行符，以便LLM更好地理解多行内容（如“主要贡献”、“奖励情况”）
                    # 同时尝试智能拆分混在一起的 KV 对
                    cell_data = [clean_cell_text(cell.text) for cell in row.cells]
                    rows.append(cell_data)
                
                # 直接转换为 DataFrame，不假设第一行是表头
                # 这样可以保留 Key-Value 类型的表格结构（如简历、登记表）
                # 同时也解决了表头包含换行符导致的问题
                df = pd.DataFrame(rows)
                tables_data.append(df)

        # 3. 写入Excel
        with pd.ExcelWriter(output_excel_path, engine='openpyxl') as writer:
            # 写入文本内容
            if text_content:
                df_text = pd.DataFrame(text_content)
                df_text.to_excel(writer, sheet_name="Text_Content", index=False)
            
            # 写入表格 (不写入 Header，不写入 Index)
            for i, df in enumerate(tables_data):
                sheet_name = f"Table_{i+1}"
                df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                
        logger.info(f"成功从 {docx_path} 提取数据到 {output_excel_path}")
        return True
    except Exception as e:
        logger.error(f"提取数据失败: {str(e)}")
        return False
